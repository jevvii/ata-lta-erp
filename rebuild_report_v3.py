#!/usr/bin/env python3
"""
Rebuild Week_1_Accomplishments_Report as v3 with ALL A-J module tables from v2,
plus proper professional formatting throughout.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "1F3864"
ALT_ROW_BG = "F2F2F2"
BODY_SIZE = Pt(12)
H1_SIZE = Pt(18)
H2_SIZE = Pt(14)
H3_SIZE = Pt(12)
TBL_HDR_SIZE = Pt(10)
TBL_BODY_SIZE = Pt(10)
SMALL_SIZE = Pt(10)


def _force_font(run, size=BODY_SIZE, bold=False, italic=False, color=BLACK):
    f = run.font
    f.name = FONT; f.size = size; f.bold = bold; f.italic = italic; f.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rF = rPr.find(qn('w:rFonts'))
    if rF is None:
        rF = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rF)
    for attr in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'):
        rF.set(qn(attr), FONT)

def R(p, text, bold=False, italic=False, size=BODY_SIZE, color=BLACK):
    run = p.add_run(text)
    _force_font(run, size, bold, italic, color)
    return run

def _shade(cell, hx):
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:shd'))
    if old is not None: tcPr.remove(old)
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hx}" w:val="clear"/>'))

def _cell_w(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:tcW'))
    if old is not None: tcPr.remove(old)
    tcPr.append(parse_xml(f'<w:tcW {nsdecls("w")} w:w="{dxa}" w:type="dxa"/>'))

def _tbl_borders(table):
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None: tblPr.remove(old)
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="A6A6A6"/>'
        '<w:left w:val="single" w:sz="6" w:space="0" w:color="A6A6A6"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="A6A6A6"/>'
        '<w:right w:val="single" w:sz="6" w:space="0" w:color="A6A6A6"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '</w:tblBorders>'))

def _tbl_w(table, inches=6.5):
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        table._tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblW'))
    if old is not None: tblPr.remove(old)
    tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="{int(inches*1440)}" w:type="dxa"/>'))

def add_h(doc, text, level=1):
    p = doc.add_paragraph()
    R(p, text, bold=True, size={1:H1_SIZE,2:H2_SIZE,3:H3_SIZE}.get(level, H2_SIZE))
    p.paragraph_format.space_before = Pt(14) if level <= 2 else Pt(10)
    p.paragraph_format.space_after = Pt(6)
    if level == 1: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_p(doc, text, bold=False, italic=False, sa=Pt(6)):
    p = doc.add_paragraph()
    R(p, text, bold=bold, italic=italic)
    p.paragraph_format.space_after = sa
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    return p

def add_b(doc, label, desc, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.left_indent = Inches(0.4 + indent*0.25)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.line_spacing = 1.15
    R(p, "\u2022  ")
    R(p, label, bold=True)
    if desc: R(p, f"  {desc}")
    return p

def make_tbl(doc, headers, rows, cw=None):
    nc = len(headers)
    t = doc.add_table(rows=1+len(rows), cols=nc)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    _tbl_borders(t); _tbl_w(t, 6.5)
    if cw is None: cw = [6.5/nc]*nc
    dxa = [int(w*1440) for w in cw]
    for ci, ht in enumerate(headers):
        c = t.rows[0].cells[ci]
        _shade(c, HEADER_BG); _cell_w(c, dxa[ci])
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.paragraphs[0].clear()
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        R(c.paragraphs[0], ht, bold=True, size=TBL_HDR_SIZE, color=WHITE)
    for ri, rd in enumerate(rows):
        bg = ALT_ROW_BG if ri%2==1 else "FFFFFF"
        for ci in range(nc):
            c = t.rows[1+ri].cells[ci]
            _shade(c, bg); _cell_w(c, dxa[ci])
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            c.paragraphs[0].clear()
            R(c.paragraphs[0], rd[ci] if ci<len(rd) else "", size=TBL_BODY_SIZE)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(2)
    sp.paragraph_format.space_after = Pt(4)
    return t


# ══════════════════════════════════════
#  BUILD
# ══════════════════════════════════════
doc = Document()
for s in doc.sections:
    s.top_margin=Inches(1); s.bottom_margin=Inches(1)
    s.left_margin=Inches(1); s.right_margin=Inches(1)

st = doc.styles['Normal']
st.font.name = FONT; st.font.size = BODY_SIZE; st.font.color.rgb = BLACK
rPr = st.element.get_or_add_rPr()
rF = rPr.find(qn('w:rFonts'))
if rF is None:
    rF = parse_xml(f'<w:rFonts {nsdecls("w")}/>'); rPr.insert(0, rF)
for a in ('w:ascii','w:hAnsi','w:eastAsia','w:cs'): rF.set(qn(a), FONT)

# ── LETTERHEAD ──
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
R(p,"MicroAxis Solutions Corp.",bold=True,size=Pt(16))
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
R(p,"C2-15 Centris Walk, Eton Centris, Quezon City, 1101",size=SMALL_SIZE)
p.paragraph_format.space_after=Pt(0)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
R(p,"Email: sales@microaxis.ph  |  Viber: +639000000000",size=SMALL_SIZE)
p.paragraph_format.space_after=Pt(6)
hr=doc.add_paragraph(); hr.alignment=WD_ALIGN_PARAGRAPH.CENTER
R(hr,"\u2500"*72,size=Pt(7),color=RGBColor(0xA0,0xA0,0xA0))
hr.paragraph_format.space_after=Pt(4)

# ── TITLE ──
add_h(doc,"WEEKLY PROGRESS REPORT",1)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
R(p,"ATA & LTA ERP Prototype \u2014 Week 1 Accomplishments",bold=True,size=Pt(13))
p.paragraph_format.space_after=Pt(4)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
R(p,"A Browser-Only Prototype Deliverable for Dual-Entity Accounting Workflow Management\nRequirements Finalization Through Rapid Prototyping",italic=True,size=Pt(11))
p.paragraph_format.space_after=Pt(10)

pm=doc.add_paragraph()
R(pm,"Prepared By:  ",bold=True,size=Pt(11)); R(pm,"Jayvee Marcelo (Full-Stack Developer / Prototype Lead)",size=Pt(11))
pm.paragraph_format.space_after=Pt(2)
pm=doc.add_paragraph()
R(pm,"Status:  ",bold=True,size=Pt(11)); R(pm,"Prototype Stage",size=Pt(11))
R(pm,"   |   ",size=Pt(11),color=RGBColor(0xA0,0xA0,0xA0))
R(pm,"Date:  ",bold=True,size=Pt(11)); R(pm,"June 6, 2026",size=Pt(11))
R(pm,"   |   ",size=Pt(11),color=RGBColor(0xA0,0xA0,0xA0))
R(pm,"Phase:  ",bold=True,size=Pt(11)); R(pm,"Requirements Finalization & Pre-Validation Complete",size=Pt(11))
pm.paragraph_format.space_after=Pt(6)
hr=doc.add_paragraph(); hr.alignment=WD_ALIGN_PARAGRAPH.CENTER
R(hr,"\u2500"*72,size=Pt(7),color=RGBColor(0xA0,0xA0,0xA0))
hr.paragraph_format.space_after=Pt(4)


# ══════════════════════════════════════
#  1) EXECUTIVE SUMMARY
# ══════════════════════════════════════
add_h(doc,"1)  Executive Summary",2)
add_p(doc,"Week 1 of the ATA & LTA ERP project has delivered a comprehensive, interactive browser-only prototype that validates the dual-entity accounting workflow, role-based access control, and all six operational stages from work request initiation through documentation handover. All 19 feedback items from ERP Feedback v1 have been analyzed and resolved. The prototype is ready for second-pass client review and final requirements sign-off before proceeding to stack selection and production development.")
add_p(doc,"This deliverable is not a production build. The current system is implemented entirely in vanilla JavaScript, HTML, and CSS with localStorage persistence. Its purpose is to replace static wireframes with interactive, data-backed screens that act as a living specification, enabling the client to validate business rules, user flows, and UI patterns before any stack or layer decisions are committed.",italic=True)


# ══════════════════════════════════════
#  2) ALIGNMENT WITH TIMELINE
# ══════════════════════════════════════
add_h(doc,"2)  Alignment with Proposed Timeline & Phases",2)
add_p(doc,"The original ERP Development Proposal (May 2026) outlines a 10\u201312 week delivery across six phases. Week 1 has completed Phases 1 and 2 in full, and has prototyped / pre-validated the core capabilities of Phases 3 and 4 as browser-level proof-of-concept. This accelerated progress means the project is on track to enter true Sprint 1 coding immediately upon client sign-off, with zero ambiguity on module scope or acceptance criteria.")

# Phase 1
add_h(doc,"Phase 1: Discovery & Planning (Proposal Week 1) \u2014 COMPLETED",3)
add_b(doc,"Key Accomplishment:","Instead of static documentation, the prototype itself is the requirements artifact. Stakeholders interact with real data, real gating rules, and real edge cases (cycle detection, role blocks, overdue aging) rather than reading about them.")

# Phase 2
add_h(doc,"Phase 2: UI/UX Design & Prototyping (Proposal Weeks 1\u20132) \u2014 COMPLETED",3)
add_b(doc,"Key Accomplishment:","The prototype went beyond wireframes by including responsive breakpoints, print media queries, empty-state patterns, and view-mode toggles (Table / Board / List). This means the Phase 3 frontend developer inherits a validated design system, not a blank canvas.")

# Phase 3
add_h(doc,"Phase 3: Core Development \u2014 Sprint 1 (Proposal Weeks 2\u20135) \u2014 PROTOTYPE PRE-VALIDATED",3)
add_p(doc,'The proposal defines Sprint 1 as: "Authentication, user/role management, multi-entity client management, and workflow engine (task creation, assignment, dependencies, recurring templates). Backend APIs and frontend screens are developed in parallel. Deliverable: Deployed alpha on staging with core modules functional."',italic=True)
add_p(doc,"The following detailed feature breakdown maps each Sprint 1 module to its prototype status and production remarks:",bold=True)

# ── TABLE A: Authentication & RBAC ──
add_p(doc,"A.  Authentication & Role Access Module",bold=True,sa=Pt(4))
make_tbl(doc, ["Features","Status","Remarks"], [
    ["\u2022 Secure login authentication","Prototype","Browser-based mock auth. To be implemented with bcrypt + JWT sessions via Node.js/Express upon workflow."],
    ["\u2022 Session management","Prototype","sessionStorage persistence. Production: Redis-backed sessions with refresh tokens."],
    ["\u2022 Password encryption","Prototype","Plaintext in localStorage. Production: bcrypt hashing with Argon2 option."],
    ["\u2022 Role-based access control","Prototype","Client-side Auth.can() matrix. Production: PostgreSQL RLS + server-side middleware RBAC."],
    ["\u2022 Multi-entity access permissions","Prototype","ATA/LTA toggle in header. Production: row-level security policies per entity."],
    ["\u2022 Admin review and approval controls","Prototype","PendingChanges gate with side-by-side diff. Production: async notification queues (SendGrid/Resend) + audit trail."],
    ["\u2022 User activity logging","Prototype","In-memory audit entries. Production: persistent audit_log table with immutable timestamps."],
    ["\u2022 Audit trail monitoring","Prototype","Admin panel audit tab. Production: SIEM-compatible export + 7-year retention."],
    ["\u2022 User Roles","Prototype",""],
    ["    \u25CB Admin","Prototype","Full system access, cross-entity visibility."],
    ["    \u25CB Manager","Prototype","Module-level approve/edit per assigned entity."],
    ["    \u25CB Staff","Prototype","Create/edit own records; submit for approval."],
    ["    \u25CB Viewer","Prototype","Read-only across all modules in assigned entity."],
], cw=[2.4, 0.9, 3.2])

# ── TABLE B: Client Management ──
add_p(doc,"B.  Client Management Module",bold=True,sa=Pt(4))
make_tbl(doc, ["Features","Status","Remarks"], [
    ["\u2022 Create client record","Prototype","localStorage CRUD. Production: PostgreSQL with unique TIN validation + duplicate detection."],
    ["\u2022 Edit client information","Prototype","Inline form with validation. Production: optimistic locking to prevent concurrent edits."],
    ["\u2022 Archive / restore clients","Prototype","Status flag toggle. Production: soft-delete with archive_reason and deleted_by fields."],
    ["\u2022 Search and filter clients","Prototype","Client-side string matching. Production: PostgreSQL full-text search (tsvector) or Meilisearch."],
    ["\u2022 Related companies (Parent, Subsidiary, Sister, Affiliate)","Prototype","Bidirectional linking in localStorage. Production: graph relationship table with CASCADE updates."],
    ["\u2022 Contact details multi-entry","Prototype","Dynamic form rows. Production: normalized contacts table (phone, email, label)."],
    ["\u2022 Point of contact assignment","Prototype","Employee dropdown. Production: foreign key constraint + assignment history log."],
    ["\u2022 Retainer flag and status","Prototype","Boolean checkbox. Production: retainer_schedule table with auto-generation cron."],
    ["\u2022 Entity-scoped visibility","Prototype","Filter by active entity. Production: RLS policy ensuring Staff never sees cross-entity data."],
], cw=[2.4, 0.9, 3.2])

# ── TABLE C: Workflow & Task Management ──
add_p(doc,"C.  Workflow & Task Management Module",bold=True,sa=Pt(4))
make_tbl(doc, ["Features","Status","Remarks"], [
    ["\u2022 Create work request","Prototype","Form with client + assignee. Production: workflow engine microservice with event sourcing."],
    ["\u2022 Edit work request details","Prototype","Modal form. Production: versioned edits with diff tracking."],
    ["\u2022 Six-stage pipeline routing","Prototype","Draft \u2192 Pre-processing \u2192 Processing \u2192 Billing \u2192 Disbursement \u2192 Completed. Production: state machine engine with rollback."],
    ["\u2022 Phase transition validation","Prototype","Client-side gating checks. Production: server-side workflow orchestrator with rollback capability."],
    ["\u2022 Create and assign tasks","Prototype","Task list with assignee dropdown. Production: task queue with notification dispatch."],
    ["\u2022 Task dependency engine (DAG)","Prototype","Predecessor selection + cycle detection. Production: topological sort service with cycle detection."],
    ["\u2022 Time logging per task","Prototype","Start/end time with 0.25h rounding. Production: PostgreSQL interval type + overtime calculation rules."],
    ["\u2022 Upload task documents (metadata)","Prototype","Filename + uploader stored. Production: AWS S3 / Cloudflare R2 presigned URLs + virus scanning."],
    ["\u2022 Retainer template creation","Prototype","Manager-only reusable templates. Production: template versioning + automated cron generation."],
    ["\u2022 Generate from retainer template","Prototype","One-click clone to new work request. Production: scheduled job runner (BullMQ / Celery)."],
    ["\u2022 Archive cancelled work requests","Prototype","Dedicated archive tab. Production: soft-delete with restore audit."],
    ["\u2022 Routing readiness badges","Prototype","Client-side blocker calculation. Production: real-time WebSocket status updates."],
], cw=[2.4, 0.9, 3.2])

# ── TABLE D: Dashboard & Analytics ──
add_p(doc,"D.  Dashboard & Analytics Preview",bold=True,sa=Pt(4))
make_tbl(doc, ["Features","Status","Remarks"], [
    ["\u2022 Bento-grid KPI widgets","Prototype","SVG mock charts. Production: React/Vue component library with real-time data fetching."],
    ["\u2022 Revenue vs. disbursement chart","Prototype","SVG bar chart from localStorage arrays. Production: Chart.js / Recharts + PostgreSQL aggregation queries."],
    ["\u2022 Entity-scoped vs. consolidated views","Prototype","Filter toggle. Production: parameterized SQL views with RLS."],
    ["\u2022 Due-date urgency warnings","Prototype","Client-side date comparison. Production: scheduled cron job for daily digest emails."],
    ["\u2022 End-of-day time log reminder","Prototype","Manila-time banner at 5 PM. Production: push notification service (FCM / OneSignal)."],
], cw=[2.4, 0.9, 3.2])

# Phase 3 Assumptions
add_p(doc,"Prototype Assumptions for Phase 3 (To Be Built in Production):",bold=True)
for lbl,desc in [
    ("Real Backend:","The prototype\u2019s localStorage DB engine will be replaced by a PostgreSQL database with row-level security policies for ATA/LTA segregation."),
    ("Real Authentication:","The mock login will be replaced by JWT-based sessions with bcrypt hashing, short expiry, and refresh rotation."),
    ("Real APIs:","All CRUD operations currently hitting localStorage will become RESTful API calls (Node.js/Django) with server-side validation."),
    ("Real Timezone Handling:","The Manila-time-aware end-of-day reminder is client-side Intl.DateTimeFormat; production will use server-side UTC storage with Asia/Manila rendering."),
    ("Board View Interactivity:","The prototype board view is read-only (no drag-and-drop). Production will implement drag-and-drop status changes with optimistic UI updates."),
    ("Risk Reduction:","Because these modules were prototyped and client-tested in Week 1, the Sprint 1 development timeline can be compressed from 3 weeks to 2\u20132.5 weeks."),
]: add_b(doc, lbl, desc, indent=1)


# Phase 4
add_h(doc,"Phase 4: Core Development \u2014 Sprint 2 (Proposal Weeks 5\u20138) \u2014 PROTOTYPE PRE-VALIDATED",3)
add_p(doc,'The proposal defines Sprint 2 as: "Billing module, disbursement module, document management system, and reporting dashboards. Integration between modules is completed and tested internally. Deliverable: Deployed beta with all MVP modules integrated."',italic=True)
add_p(doc,"The following detailed feature breakdown maps each Sprint 2 module to its prototype status and production remarks:",bold=True)

# ── TABLE E: Billing Module ──
add_p(doc,"E.  Billing Module",bold=True,sa=Pt(4))
make_tbl(doc, ["Features","Status","Remarks"], [
    ["\u2022 Create sales invoice","Prototype","PF + Gov Fee line items. Production: atomic invoice numbering (ATA-SI-YYYY-NNN sequence) + tax compliance."],
    ["\u2022 Edit invoice details","Prototype","Line item CRUD. Production: draft \u2192 sent immutability rules + edit audit trail."],
    ["\u2022 Record payment","Prototype","Cash/Check/Bank Transfer/GCash/Maya. Production: payment gateway integration (Stripe / PayMongo)."],
    ["\u2022 BIR-compliant invoice PDF","Prototype","Browser @media print CSS. Production: Puppeteer / WeasyPrint server-side with BIR e-invoicing compliance."],
    ["\u2022 Voucher print (no-header receipt)","Prototype","Stripped print CSS. Production: same PDF engine with voucher template."],
    ["\u2022 Aging report (0-30, 31-60, 61-90, 90+)","Prototype","Client-side date bucketing. Production: materialized view with indexed due_date."],
    ["\u2022 Trash and restore invoices","Prototype","Cancelled status flag. Production: soft-delete with trash_retention_days policy."],
    ["\u2022 Recurring billing templates","Prototype","Manual generate button. Production: cron job with invoice auto-generation and email dispatch."],
    ["\u2022 Payment history tracking","Prototype","Inline table. Production: payments table with foreign key + reconciliation status."],
    ["\u2022 Outstanding balance calculation","Prototype","Client-side reduce(). Production: SQL SUM() with index on invoice_id."],
], cw=[2.4, 0.9, 3.2])

# ── TABLE F: Disbursement & Expense ──
add_p(doc,"F.  Disbursement & Expense Module",bold=True,sa=Pt(4))
make_tbl(doc, ["Features","Status","Remarks"], [
    ["\u2022 File expense","Prototype","Category + amount + receipt metadata. Production: receipt OCR (Tesseract / AWS Textract) for auto-fill."],
    ["\u2022 Two-tier approval chain","Prototype","Manager \u2192 Accounting for Firm Fund. Production: approval workflow engine with escalation rules."],
    ["\u2022 One-tier approval (Client Fund)","Prototype","Accounting direct verify. Production: same workflow engine with conditional path logic."],
    ["\u2022 Payment handler tracking","Prototype","Requested By + Handled By fields. Production: foreign key to users with release authorization."],
    ["\u2022 Generate expense PDF","Prototype","Browser print. Production: server-side PDF with digital signature (DocuSign / local PKI)."],
    ["\u2022 Generate disbursement voucher PDF","Prototype","Structured print layout. Production: voucher template engine with barcode/QR."],
    ["\u2022 Summary report by category/employee","Prototype","Client-side grouping. Production: OLAP-style aggregations with date range indexing."],
    ["\u2022 Recurring disbursement templates","Prototype","Manual generate button. Production: cron-scheduled auto-generation with approval bypass for retainers."],
    ["\u2022 Pending for release section","Prototype","Filter for Approved status + handler match. Production: real-time notification on status change."],
], cw=[2.4, 0.9, 3.2])

# ── TABLE G: DMS ──
add_p(doc,"G.  Document Management System (DMS)",bold=True,sa=Pt(4))
make_tbl(doc, ["Features","Status","Remarks"], [
    ["\u2022 Upload document metadata","Prototype","Filename + category + uploader. Production: multipart upload to S3/R2 with virus scan + checksum."],
    ["\u2022 Document lifecycle transitions","Prototype","collected \u2192 with_documentations \u2192 scanned \u2192 in_envelope \u2192 stored. Production: state machine with audit log."],
    ["\u2022 Handover log (chain of custody)","Prototype","Text entries. Production: immutable handover_log table with digital signature."],
    ["\u2022 Envelope labeling","Prototype","Auto-generated text ENVELOPE-[clientId]-[seq]. Production: barcode label generation (Zebra API)."],
    ["\u2022 Admin comments","Prototype","Comment thread in detail view. Production: comment table with @mentions and notification dispatch."],
    ["\u2022 Entity-scoped document filtering","Prototype","Filter by active entity. Production: RLS policy on documents table."],
    ["\u2022 View mode toggles (Table/Board/List)","Prototype","Client-side rendering. Production: server-side pagination + infinite scroll."],
], cw=[2.4, 0.9, 3.2])

# ── TABLE H: Transmittal ──
add_p(doc,"H.  Transmittal Module",bold=True,sa=Pt(4))
make_tbl(doc, ["Features","Status","Remarks"], [
    ["\u2022 Create transmittal letter","Prototype","Select DMS documents + tracking number. Production: auto-sequential numbering with collision guard."],
    ["\u2022 Edit transmittal details","Prototype","Itemized document list CRUD. Production: versioned transmittal with change audit."],
    ["\u2022 Generate tracking number","Prototype","ATA-TX-YYYY-NNN format. Production: database sequence with entity prefix + year reset."],
    ["\u2022 Generate transmittal letter PDF","Prototype","Browser print with entity header. Production: server-side PDF with digital signature block."],
    ["\u2022 Acknowledge receipt","Prototype","Received By + date text fields. Production: e-signature capture (canvas / DocuSign) + timestamped log."],
], cw=[2.4, 0.9, 3.2])

# ── TABLE I: Reports & Analytics ──
add_p(doc,"I.  Reports & Analytics Module",bold=True,sa=Pt(4))
make_tbl(doc, ["Features","Status","Remarks"], [
    ["\u2022 Analytics dashboard (bento-grid)","Prototype","Mock data KPIs. Production: real-time data warehouse (PostgreSQL materialized views or ClickHouse)."],
    ["\u2022 Daily task report","Prototype","Today\u2019s completions + time logs. Production: scheduled daily report generation + email delivery."],
    ["\u2022 Weekly performance summary","Prototype","Completed / pending / overdue counts. Production: trend analysis with week-over-week comparison."],
    ["\u2022 Monthly pending tasks list","Prototype","Pending + recurring retainer tasks. Production: automated monthly digest with SLA alerting."],
    ["\u2022 Entity P&L snapshot","Prototype","Revenue vs. disbursement bar chart. Production: financial reporting module with drill-down."],
], cw=[2.4, 0.9, 3.2])

# ── TABLE J: Admin / Users ──
add_p(doc,"J.  Admin / Users Module",bold=True,sa=Pt(4))
make_tbl(doc, ["Features","Status","Remarks"], [
    ["\u2022 Create user account","Prototype","Form with role + entity assignment. Production: invite email flow (SendGrid) + temporary password."],
    ["\u2022 Edit user account","Prototype","Role/entity update. Production: role change audit + re-auth requirement on role escalation."],
    ["\u2022 Activate / deactivate users","Prototype","Status toggle. Production: soft-delete with deactivation_reason and reactivation log."],
    ["\u2022 View audit log","Prototype","Read-only action history. Production: tamper-proof audit table with hash chain verification."],
    ["\u2022 Review pending approvals","Prototype","Side-by-side diff. Production: optimistic locking + concurrent approval conflict resolution."],
    ["\u2022 Approve / reject pending changes","Prototype","Instant state update. Production: async webhook on approval + downstream cascade triggers."],
    ["\u2022 Self-approval block","Prototype","Auth.isSelfApprover() check. Production: server-side enforcement with immutable audit entry."],
    ["\u2022 Reset data to seed","Prototype","One-click localStorage reset. Production: admin-only CLI command with backup requirement."],
], cw=[2.4, 0.9, 3.2])

# Phase 4 Assumptions
add_p(doc,"Prototype Assumptions for Phase 4 (To Be Built in Production):",bold=True)
for lbl,desc in [
    ("Real File Storage:","File uploads currently capture metadata only due to localStorage 5MB limits. Production will use AWS S3 / Cloudflare R2."),
    ("Real PDF Generation:","Current PDF output uses browser @media print CSS. Production will use a server-side PDF library."),
    ("Real Invoice Numbering:","Sequential entity-specific numbering is simulated. Production will use database sequences with atomic increments."),
    ("Real Approval Workflows:","Approval chains are instant state changes. Production will implement async notification queues and audit trails."),
    ("Real Reporting Engine:","Reports aggregate localStorage arrays in memory. Production will use SQL aggregations with indexed queries."),
    ("Automated Recurring Generation:","Templates require manual clicks. Production will implement server-side cron jobs."),
    ("Full-Text Search:","DMS search is metadata-only. Production will add PostgreSQL tsvector or Meilisearch."),
    ("Risk Reduction:","Sprint 2 timeline can be compressed from 3 weeks to 2\u20132.5 weeks with additional time for integration testing."),
]: add_b(doc, lbl, desc, indent=1)

# Phase 5
add_h(doc,"Phase 5: Integration Testing & UAT (Proposal Weeks 8\u201310) \u2014 PLANNED",3)
add_b(doc,"Prototype Contribution to Phase 5:","The 19 resolved feedback items from ERP Feedback v1 constitute an early UAT cycle. Issues like BIR PDF formatting, payment history overflow, and role clarity were caught and fixed before a single production commit.")

# Phase 6
add_h(doc,"Phase 6: Go-Live & Post-Launch Support (Proposal Weeks 10\u201312) \u2014 PLANNED",3)
add_p(doc,"No prototype deliverables for this phase. Go-live activities begin after production beta passes UAT.")


# ══════════════════════════════════════
#  3) MODULE BREAKDOWN (bullet format)
# ══════════════════════════════════════
add_h(doc,"3)  Week 1 Prototype Accomplishments \u2014 Module Breakdown",2)

for title, items in [
    ("A)  Foundation & Data Layer (Browser-Only)", [
        ("localStorage DB Engine.","Full CRUD persistence layer with schema versioning (v3), automatic migration from v2, and seed data hydration on first load."),
        ("Multi-Entity Segregation.","Every record is tagged to ATA or LTA. The entity switcher changes the active context and re-renders all views atomically."),
        ("Authentication & RBAC.","Login with email/password, session restoration, four-tier role gates (Admin/Manager/Staff/Viewer), and permission matrix."),
        ("Schema v3 Design.","Tables: users, clients, workRequests, tasks, invoices, disbursements, documents, transmittals, retainerTemplates, billingTemplates, disbursementTemplates, pendingChanges."),
    ]),
    ("B)  Dashboard & Analytics (Preview)", [
        ("Bento-Grid KPI Widgets.","Active Work Requests, Total Clients, Disbursement Summary, and Revenue vs. Disbursement bar chart."),
        ("Entity-Scoped vs. Consolidated Views.","Admin sees combined ATA + LTA data; non-admin users see only their assigned entity."),
        ("Due-Date Warnings.","Work Requests due this week are highlighted with color-coded urgency."),
    ]),
    ("C)  Client Management Module", [
        ("Full CRUD.","Create, edit, and archive clients with mandatory entity assignment, TIN, trade name, and business address."),
        ("Related Companies.","Multi-entry UI for Parent, Subsidiary, Sister Company, and Affiliate relationships."),
        ("Contact Details.","Repeatable rows for mobile, landline, email, and alternate contacts with labels."),
        ("Point of Contact.","Employee selector assigning a staff member as the primary liaison."),
        ("Retainer Flag.","Checkbox marking the client as a recurring retainer; enables monthly/quarterly template generation."),
    ]),
    ("D)  Operations (Workflow & Task Management)", [
        ("Six-Stage Pipeline.","Draft \u2192 Pre-processing \u2192 Processing \u2192 Billing \u2192 Disbursement \u2192 Completed."),
        ("Task Dependency Engine (DAG).","Predecessor selection per task; cycle detection prevents circular dependencies."),
        ("Time Logging.","Retrospective entry with start time, end time, auto-calculated hours (rounded to 0.25)."),
        ("Document Upload.","File metadata capture stored in taskDocuments[]."),
        ("Retainer Templates.","Manager-only feature: create reusable task templates with monthly or quarterly schedules."),
        ("Archive for Cancelled Work Requests.","Cancelled items moved to a dedicated Archive tab. Restore to Draft available."),
    ]),
    ("E)  Billing Module", [
        ("Sales Invoice Creation.","Line items for Professional Fee and Government Fee; auto-calculated subtotal."),
        ("Payment Recording.","Manual entry of amount, method, reference number, date, and collected-by employee."),
        ("Aging Report.","Overdue buckets: 0\u201330, 31\u201360, 61\u201390, 90+ days with color-coded totals."),
        ("BIR-Compliant Invoice PDF.","Browser print with seller header, buyer info, itemized list, totals, and BIR footer."),
        ("Voucher Print.","Stripped seller header for pass-through government fee receipts."),
        ("Trash & Restore.","Draft invoices can be moved to Trash; restored back to Draft."),
        ("Recurring Billing Templates.","Manual \"Generate Next Period\" button clones template into a new invoice."),
    ]),
    ("F)  Disbursement & Expense Module", [
        ("Expense Filing.","Staff submits category, description, amount, receipt metadata, fund source, and linked work request."),
        ("Approval Chain.","1-tier or 2-tier workflow: Submitted \u2192 Under Review \u2192 Approved \u2192 Released \u2192 Rejected."),
        ("Payment Handler Tracking.","Explicit \"Requested By\" and \"Payment Handled By\" fields."),
        ("Disbursement PDF / Voucher.","Structured print with Voucher No, Date, Payee, Payment Method, Description, Amount, and signature lines."),
        ("Recurring Templates.","New Template form with Name, Category, Amount, Fund Source, and Schedule."),
    ]),
    ("G)  Document Management System (DMS)", [
        ("File Upload.","Per-work-request upload with 4 categories."),
        ("Document Lifecycle.","Status transitions: collected \u2192 with_documentations \u2192 scanned \u2192 in_envelope \u2192 stored."),
        ("Handover Log.","Chain-of-custody entries for physical original copies handed to clients."),
        ("Envelope Labeling.","Auto-generated text labels: ENVELOPE-[clientId]-[sequence]."),
        ("Admin Comments.","Comment threads on document detail view; restricted to Admin role."),
    ]),
    ("H)  Transmittal Module", [
        ("Tracking Number Generation.","Auto-sequential per entity: ATA-TX-2025-001, LTA-TX-2025-001."),
        ("Itemized Document List.","Add/remove rows linking existing DMS documents."),
        ("Acknowledgment.","\"Received By\" text field + date entry when client signs the transmittal letter."),
        ("Transmittal Letter PDF.","Formal print with entity header, itemized list, signature block, and tracking number."),
    ]),
    ("I)  Reports Module (MVP-lite Preview)", [
        ("Daily Report.","Tasks completed today per employee with start/end times and total hours."),
        ("Weekly Summary.","Completed vs. pending vs. overdue per employee; overdue warnings highlighted."),
        ("Monthly Pending List.","Pending tasks grouped by employee plus recurring retainer tasks due this month."),
    ]),
    ("J)  Admin / Users Module", [
        ("User CRUD.","Create, edit, activate, and deactivate employee accounts with role assignment and entity scoping."),
        ("Pending Approvals (Admin Review Gate).","Side-by-side diff of proposed changes vs. current record; Approve/Reject actions."),
        ("Self-Approval Block.","Users cannot approve their own disbursements or review their own pending changes."),
    ]),
]:
    add_h(doc, title, 3)
    for lbl, desc in items:
        add_b(doc, lbl, desc)


# ══════════════════════════════════════
#  4) CLIENT FEEDBACK
# ══════════════════════════════════════
add_h(doc,"4)  Client Feedback Integration \u2014 ERP Feedback v1",2)
add_p(doc,"The Week 1 prototype was delivered to the client for first-pass review on May 2026. The client returned 19 structured feedback items. All items have been analyzed, triaged, and resolved in the prototype.")

for title, items in [
    ("A)  Global UI & Layout Fixes", [
        ("Dashboard Widget Consistency.","Issue #1: WR Due for Week missing from consolidated view. Resolved."),
        ("Save/Cancel Button Placement.","Issue #2: Buttons at bottom. Resolved by moving to top-right form-header-bar pattern."),
        ("Board View Horizontal Scroll.","Issues #9, #10: Resolved via proportional flex columns and word-wrap."),
        ("View-Mode Toggle Placement.","Issue #12: Relocated toggle below the filters bar."),
    ]),
    ("B)  Admin Review Gate Fixes", [
        ("Rejected Submissions Visibility.","Issue #3: Added Rejected section with rejection reason and Resubmit button."),
    ]),
    ("C)  Clients Module Fixes", [
        ("Missing Table Columns.","Issue #4: Added Related Companies and Contact Details columns."),
    ]),
    ("D)  Billing Module Fixes", [
        ("Invoice PDF Format.","Issue #5: Added BIR-compliant seller header, buyer info, line items, and footer."),
        ("Payment Details in List Views.","Issue #6: Added Paid and Balance columns to table view."),
        ("Voucher Print CSS.","Issue #7: Resolved with @media print CSS."),
        ("Payment History Overflow.","Issue #8: Wrapped in overflow-x:auto container."),
    ]),
    ("E)  Disbursement Module Fixes", [
        ("Role Clarity.","Issue #11: Added explicit \"Requested By\" and \"Payment Handled By\" labels."),
        ("PDF Generation.","Issue #13: Structured disbursement PDF with signature lines."),
        ("Recurring Templates.","Issue #14: Added New Template form with Generate button."),
        ("Payment Info in List Views.","Issue #6 (Disbursement part): Added Payment Method and Handled By columns."),
    ]),
    ("F)  Operations Module Fixes", [
        ("Inline Time Logs.","Issue #15: Added expandable accordion rows per task."),
        ("Documentation Staff Visibility.","Issue #16: Added Auth.can('dms:handover') exception."),
        ("Month Filter Reliability.","Issue #17: Replaced with select dropdown of 12 months + year."),
        ("Admin File Access.","Issue #18: Added clickable links to DMS records."),
        ("Inline Task Documents.","Issue #19: Embedded document metadata inside collapsible rows."),
    ]),
    ("G)  Additional Enhancements Beyond Feedback v1", [
        ("Widely Compatible View Icons.","Centralized ViewIcons utility (Lucide-style SVGs) for cross-browser compatibility."),
        ("Billing \u2192 Disbursement Routing Fix.","Relaxed disbursement linkage gate for task-level and work-request-level links."),
        ("Actionable Dependency Hints.","Routing blocker messages with Go button to target module."),
        ("End-of-Day Time Log Reminder.","Manila-time-aware banner at 5:00 PM+."),
        ("Routing Readiness Badges.","\"Ready to route\" (green) or \"N pending\" (orange) badges with tooltips."),
        ("Status Color Consistency.","Unified phase colors across all UI elements."),
    ]),
]:
    add_h(doc, title, 3)
    for lbl, desc in items:
        add_b(doc, lbl, desc)


# ══════════════════════════════════════
#  5) KNOWN LIMITATIONS
# ══════════════════════════════════════
add_h(doc,"5)  Known Limitations & Prototype Boundaries",2)
add_p(doc,"The following limitations are intentional for the requirements-finalization prototype. They will be addressed once the client approves the finalized specification.")
for lbl,desc in [
    ("No Server or Backend.","All data is stored in browser localStorage. No REST API, no database, no server-side validation."),
    ("No Real File Storage.","File uploads capture metadata only. Future stack will use AWS S3 / Cloudflare R2."),
    ("No Email or Notification Service.","No SendGrid, Resend, or SMTP integration."),
    ("No Mobile Native App.","Web-responsive only. Native iOS/Android is out of MVP scope."),
    ("No Third-Party Integrations.","No BIR eFPS, QuickBooks, Xero, or bank API connections."),
    ("Board View is Read-Only.","No drag-and-drop; status changes use dropdown/edit flow."),
    ("No Automated Recurring Generation.","Templates require manual clicks. No server-side cron."),
]: add_b(doc, lbl, desc)


# ══════════════════════════════════════
#  6) DELIVERABLES & FILE INVENTORY
# ══════════════════════════════════════
add_h(doc,"6)  Deliverables & File Inventory",2)
add_p(doc,"The Week 1 prototype is delivered as a self-contained set of static files. No build step, no npm install, and no server is required.")
for fn,desc in [
    ("index.html","App shell, navigation structure, login screen, module script loading order."),
    ("css/styles.css","Design system with CSS custom properties, responsive breakpoints, print media queries."),
    ("js/utils.js","DOM builder (el()), formatting, debounce, deepClone, generateId(), field helpers."),
    ("js/data.js","Schema v3 definition, seed data for all 12 tables, migration logic, DB persistence engine."),
    ("js/auth.js","Login/logout, session restore, Auth.can() permission matrix, entity switching."),
    ("js/app.js","Hash router, module loader, sidebar highlight, responsive menu toggle."),
    ("js/dashboard.js","KPI widgets, revenue chart, due-date warnings, entity-scoped rendering."),
    ("js/clients.js","Client CRUD, related companies, contact details, point-of-contact, retainer flag."),
    ("js/workflow.js","Work request engine, 6-stage routing, task DAG, time logging, retainer templates, archive."),
    ("js/billing.js","Invoice CRUD, line items, payment recording, aging report, BIR PDF, voucher print."),
    ("js/disbursement.js","Expense filing, approval chain, payment handler, disbursement PDF, recurring templates."),
    ("js/dms.js","Document upload metadata, lifecycle transitions, handover log, envelope labeling."),
    ("js/transmittal.js","Tracking number generation, itemized document list, acknowledgment, transmittal PDF."),
    ("js/reports.js","Daily report, weekly summary, monthly pending list, filterable views."),
    ("js/users.js","User CRUD, role/entity assignment, pending approvals, self-approval block."),
    ("js/pendingChanges.js","Admin Review Gate: submitForReview, approve, reject, resubmit, buildDiff."),
    ("js/datepicker.js + js/timepicker.js","Material Design-inspired date and time picker components."),
]: add_b(doc, fn+"  \u2014", desc)


# ══════════════════════════════════════
#  7) NEXT STEPS
# ══════════════════════════════════════
add_h(doc,"7)  Next Steps \u2014 Week 2 Plan",2)
add_p(doc,"Week 2 continues the requirements-finalization prototype cycle. The goal is to produce a second prototype iteration (v1.1) ready for final client sign-off before stack selection.")
for lbl,desc in [
    ("Client Feedback v1.1 Review.","Address any new or unresolved feedback items."),
    ("UI Polish Pass.","Refine board card density, table row hover states, empty-state illustrations."),
    ("Data Export / Import.","Add JSON export/import for all tables."),
    ("Stack Decision Brief.","Draft the technical architecture document based on validated requirements."),
    ("UAT Test Script Draft.","Begin writing the User Acceptance Testing checklist for Phase 5."),
]: add_b(doc, lbl, desc)


# ══════════════════════════════════════
#  8) CONCLUSION & SIGN-OFF
# ══════════════════════════════════════
add_h(doc,"8)  Conclusion & Sign-off",2)
add_p(doc,"Week 1 has successfully delivered a comprehensive, interactive browser-only prototype that validates the dual-entity accounting workflow, role-based access control, and all six operational stages from work request initiation through documentation handover.")

add_p(doc,"Proposal Deliverable Compliance \u2014 Phase 1:",bold=True,sa=Pt(4))
make_tbl(doc, ["Proposal Deliverable","Status","How the Prototype Satisfies It"], [
    ["Requirements Specification","Complete","12 modules, 80+ screens, and 19 resolved feedback items constitute a validated, living requirements spec"],
    ["Entity-Relationship Diagram (ERD)","Complete","Schema v3 designed with 12 tables: users, clients, workRequests, tasks, invoices, disbursements, documents, transmittals, retainerTemplates, billingTemplates, disbursementTemplates, pendingChanges"],
    ["Sprint Backlog","Complete","Each prototype screen maps to a user story; acceptance criteria are executable"],
    ["Business Rule Validation","Complete","Domain rules (PF vs gov't fee, entity segregation, 2-tier approval, self-approval block) are enforced in code"],
], cw=[1.7,1.0,3.8])

add_p(doc,"Proposal Deliverable Compliance \u2014 Phase 2:",bold=True,sa=Pt(4))
make_tbl(doc, ["Proposal Deliverable","Status","How the Prototype Satisfies It"], [
    ["Clickable Prototype","Complete","Fully navigable 12-module application with login, routing, modal flows, and print previews"],
    ["UI Style Guide","Complete","CSS custom properties define the design system: color tokens, typography, spacing, form patterns"],
    ["Design Approval Readiness","Complete","Client received the prototype, submitted 19 feedback items, and all items were triaged and resolved"],
], cw=[1.7,1.0,3.8])

add_p(doc,"Testing & Quality Assurance:",bold=True,sa=Pt(4))
make_tbl(doc, ["Proposal Deliverable","Status","Notes"], [
    ["Unit Testing (\u2265 80% coverage)","Planned","Will begin in Sprint 1 (Week 2); prototype acceptance criteria serve as test cases"],
    ["Integration Testing","Planned","Critical workflow end-to-end tests: WR creation \u2192 task assignment \u2192 billing \u2192 document upload"],
    ["UAT Sign-off","Planned","UAT test script draft is a Week 2 deliverable; Domain Expert (CPA/MBA) will validate financial outputs"],
    ["Security Testing","Planned","OWASP ZAP baseline scan, RBAC policy validation, dependency vulnerability scanning"],
], cw=[1.7,1.0,3.8])

add_p(doc,"Go-Live & Support:",bold=True,sa=Pt(4))
make_tbl(doc, ["Proposal Deliverable","Status","Notes"], [
    ["Production Deployment","Planned","Hosting platform (Vercel / Railway / Render) to be selected after client sign-off"],
    ["User Training","Planned","The prototype itself serves as a training sandbox"],
    ["Hypercare (14 days)","Planned","Critical bug resolution within 12 hours; non-critical within 48 hours"],
    ["Documentation","Planned","System administrator manual and end-user quick-start guide"],
], cw=[1.7,1.0,3.8])

add_p(doc,"Phase Status Summary:",bold=True,sa=Pt(4))
make_tbl(doc, ["Phase","Proposal Weeks","Status"], [
    ["Phase 1: Discovery & Planning","Week 1","COMPLETED \u2014 Requirements spec, ERD, and sprint backlog validated through prototype"],
    ["Phase 2: UI/UX Design & Prototyping","Weeks 1\u20132","COMPLETED \u2014 Clickable prototype delivered; UI style guide established; 19 feedback items resolved"],
    ["Phase 3: Sprint 1 \u2014 Core Engine","Weeks 2\u20135","PROTOTYPE PRE-VALIDATED \u2014 Auth, roles, client mgmt, workflow engine demonstrated"],
    ["Phase 4: Sprint 2 \u2014 Business Modules","Weeks 5\u20138","PROTOTYPE PRE-VALIDATED \u2014 Billing, disbursement, DMS, reporting demonstrated"],
    ["Phase 5: Testing & UAT","Weeks 8\u201310","PLANNED \u2014 UAT script draft in Week 2; formal UAT after production beta"],
    ["Phase 6: Go-Live & Support","Weeks 10\u201312","PLANNED \u2014 Deployment, training, and 14-day hypercare"],
], cw=[2.0,1.2,3.3])

add_p(doc,"Project Timeline Impact:",bold=True,sa=Pt(4))
add_p(doc,"Because Phases 1\u20132 are complete and Phases 3\u20134 capabilities are pre-validated through the prototype, the project is positioned to begin production coding immediately upon client sign-off, with compressed sprint timelines and significantly reduced scope-ambiguity risk.")
add_p(doc,"This report is submitted for client review and written approval. Upon sign-off, the project will transition from the Requirements Finalization phase to Month 1 (Milestone 1) \u2014 Platform Foundation production development.",italic=True,sa=Pt(20))

# ── SIGN-OFF ──
hr=doc.add_paragraph(); hr.alignment=WD_ALIGN_PARAGRAPH.CENTER
R(hr,"\u2500"*72,size=Pt(7),color=RGBColor(0xA0,0xA0,0xA0))
hr.paragraph_format.space_after=Pt(8)

st = doc.add_table(rows=3, cols=2)
st.alignment = WD_TABLE_ALIGNMENT.CENTER
_tbl_borders(st); _tbl_w(st, 6.5)
for row in st.rows:
    _cell_w(row.cells[0], int(3.25*1440))
    _cell_w(row.cells[1], int(3.25*1440))
for ci,txt in enumerate(["Prepared By:","Reviewed By:"]):
    c=st.rows[0].cells[ci]; _shade(c,"F2F2F2")
    c.paragraphs[0].clear(); c.paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    R(c.paragraphs[0],txt,bold=True,size=Pt(11))
cl=st.rows[1].cells[0]; _shade(cl,"FFFFFF"); cl.paragraphs[0].clear()
R(cl.paragraphs[0],"Jayvee Marcelo",bold=True,size=Pt(11))
R(cl.add_paragraph(),"Full-Stack Developer / Prototype Lead",size=SMALL_SIZE)
R(cl.add_paragraph(),"MicroAxis Solutions Corp.",size=SMALL_SIZE)
cr=st.rows[1].cells[1]; _shade(cr,"FFFFFF"); cr.paragraphs[0].clear()
cr.paragraphs[0].paragraph_format.space_before=Pt(20)
R(cr.paragraphs[0],"_"*45,size=Pt(11))
R(cr.add_paragraph(),"Client Representative",size=SMALL_SIZE)
R(cr.add_paragraph(),"Date: "+"_"*38,size=SMALL_SIZE)
dl=st.rows[2].cells[0]; _shade(dl,"FFFFFF"); dl.paragraphs[0].clear()
R(dl.paragraphs[0],"Date: June 6, 2026",size=SMALL_SIZE)
dr=st.rows[2].cells[1]; _shade(dr,"FFFFFF"); dr.paragraphs[0].clear()

doc.add_paragraph()
pe=doc.add_paragraph(); pe.alignment=WD_ALIGN_PARAGRAPH.CENTER
R(pe,"\u2014 End of Report \u2014",bold=True,size=Pt(12))

out = "Week_1_Accomplishments_Report_v3.docx"
doc.save(out)
print(f"\u2705 Saved: {out}")
