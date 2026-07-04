#!/usr/bin/env python3
"""
Rebuild Week_1_Accomplishments_Report.docx with professional formatting.
- Times New Roman throughout
- Black text
- Proper table borders and formatting
- Clean heading hierarchy
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

# ── Constants ──
FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "1F3864"       # Dark navy for table headers
ALT_ROW_BG = "F2F2F2"      # Light gray alternating rows
GREEN_TEXT = RGBColor(0x0B, 0x6E, 0x23)   # Dark green for ✅
ORANGE_TEXT = RGBColor(0xC2, 0x7B, 0x0A)  # Dark orange for 🔄
GRAY_TEXT = RGBColor(0x4B, 0x4B, 0x4B)    # Gray for ⏳

BODY_SIZE = Pt(12)
H1_SIZE = Pt(18)
H2_SIZE = Pt(14)
H3_SIZE = Pt(12)
TABLE_HEADER_SIZE = Pt(11)
TABLE_BODY_SIZE = Pt(10.5)
SMALL_SIZE = Pt(10)


# ── Helpers ──
def set_cell_shading(cell, color_hex):
    """Set the background shading of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table):
    """Apply full borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '</w:tblBorders>'
    )
    # Remove existing borders if any
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def add_run(paragraph, text, bold=False, italic=False, size=BODY_SIZE, color=BLACK, font_name=FONT):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    # Set East Asian font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    return run


def add_heading(doc, text, level=1):
    """Add a heading with Times New Roman formatting."""
    h = doc.add_heading(level=level)
    sizes = {1: H1_SIZE, 2: H2_SIZE, 3: H3_SIZE}
    add_run(h, text, bold=True, size=sizes.get(level, H2_SIZE))
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    if level == 1:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return h


def add_body(doc, text, bold=False, italic=False, indent=False, space_after=Pt(4)):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    add_run(p, text, bold=bold, italic=italic)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Inches(0.5)
    return p


def add_bullet(doc, label, description, indent_level=0):
    """Add a bullet point with bold label and normal description."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Inches(0.35 + indent_level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    add_run(p, "•  ", size=BODY_SIZE)
    add_run(p, label, bold=True, size=BODY_SIZE)
    if description:
        add_run(p, f"  {description}", size=BODY_SIZE)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Add a professionally formatted table."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set borders
    set_table_borders(table)

    # Header row
    hdr_row = table.rows[0]
    for ci, header_text in enumerate(headers):
        cell = hdr_row.cells[ci]
        set_cell_shading(cell, HEADER_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, header_text, bold=True, size=TABLE_HEADER_SIZE, color=WHITE)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[1 + ri]
        if ri % 2 == 1:
            bg = ALT_ROW_BG
        else:
            bg = "FFFFFF"

        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_shading(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]

            # Determine text color based on content
            txt_color = BLACK
            txt_bold = False
            if cell_text.startswith("✅"):
                txt_color = GREEN_TEXT
                txt_bold = True
            elif cell_text.startswith("🔄"):
                txt_color = ORANGE_TEXT
                txt_bold = True
            elif cell_text.startswith("⏳"):
                txt_color = GRAY_TEXT
                txt_bold = True

            # Center the status column (usually column index 1)
            if ci == 1 and (cell_text.startswith("✅") or cell_text.startswith("🔄") or cell_text.startswith("⏳")):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            add_run(p, cell_text, bold=txt_bold, size=TABLE_BODY_SIZE, color=txt_color)

    # Set column widths if provided
    if col_widths:
        for ri_idx in range(len(table.rows)):
            for ci_idx, width in enumerate(col_widths):
                table.rows[ri_idx].cells[ci_idx].width = width

    doc.add_paragraph()  # spacing after table
    return table


# ══════════════════════════════════════════════
#  BUILD DOCUMENT
# ══════════════════════════════════════════════

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# ── Set default font ──
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = BODY_SIZE
style.font.color.rgb = BLACK
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15

# Fix heading styles to use Times New Roman
for lvl in range(1, 5):
    try:
        hs = doc.styles[f'Heading {lvl}']
        hs.font.name = FONT
        hs.font.color.rgb = BLACK
        # Remove any color theme
        rPr = hs.element.find(qn('w:rPr'))
        if rPr is not None:
            color_el = rPr.find(qn('w:color'))
            if color_el is not None:
                color_el.set(qn('w:val'), '000000')
                if color_el.get(qn('w:themeColor')):
                    del color_el.attrib[qn('w:themeColor')]
                if color_el.get(qn('w:themeShade')):
                    del color_el.attrib[qn('w:themeShade')]
    except:
        pass


# ════════════════════════════════════════
#  HEADER BLOCK
# ════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "MicroAxis Solutions Corp.", bold=True, size=Pt(14))
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p2, "C2-15 Centris Walk, Eton Centris, Quezon City, 1101", size=SMALL_SIZE)
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p3, "Email: sales@microaxis.ph  |  Viber: +639000000000", size=SMALL_SIZE)
p3.paragraph_format.space_after = Pt(8)

# Divider line
p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_line, "─" * 75, size=Pt(8), color=RGBColor(0xBF, 0xBF, 0xBF))

# Title
add_heading(doc, "WEEKLY PROGRESS REPORT", level=1)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_sub, "ATA & LTA ERP Prototype — Week 1 Accomplishments", bold=True, size=Pt(13))
p_sub.paragraph_format.space_after = Pt(4)

p_desc = doc.add_paragraph()
p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_desc, "A Browser-Only Prototype Deliverable for Dual-Entity Accounting Workflow Management", italic=True, size=Pt(11))
p_desc2 = doc.add_paragraph()
p_desc2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_desc2, "Requirements Finalization Through Rapid Prototyping", italic=True, size=Pt(11))
p_desc2.paragraph_format.space_after = Pt(12)

# Meta info
p_meta = doc.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
add_run(p_meta, "Prepared By: ", bold=True, size=Pt(11))
add_run(p_meta, "Jayvee Marcelo (Full-Stack Developer / Prototype Lead)", size=Pt(11))

p_meta2 = doc.add_paragraph()
add_run(p_meta2, "Status: ", bold=True, size=Pt(11))
add_run(p_meta2, "Prototype Stage", size=Pt(11))
add_run(p_meta2, "    |    ", size=Pt(11), color=RGBColor(0xBF, 0xBF, 0xBF))
add_run(p_meta2, "Date: ", bold=True, size=Pt(11))
add_run(p_meta2, "June 6, 2026", size=Pt(11))
add_run(p_meta2, "    |    ", size=Pt(11), color=RGBColor(0xBF, 0xBF, 0xBF))
add_run(p_meta2, "Phase: ", bold=True, size=Pt(11))
add_run(p_meta2, "Requirements Finalization & Pre-Validation Complete", size=Pt(11))
p_meta2.paragraph_format.space_after = Pt(12)

# Divider
p_line2 = doc.add_paragraph()
p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_line2, "─" * 75, size=Pt(8), color=RGBColor(0xBF, 0xBF, 0xBF))
p_line2.paragraph_format.space_after = Pt(6)


# ════════════════════════════════════════
#  SECTION 1: EXECUTIVE SUMMARY
# ════════════════════════════════════════
add_heading(doc, "1)  Executive Summary", level=2)

add_body(doc, "Week 1 of the ATA & LTA ERP project has delivered a comprehensive, interactive browser-only prototype that validates the dual-entity accounting workflow, role-based access control, and all six operational stages from work request initiation through documentation handover. All 19 feedback items from ERP Feedback v1 have been analyzed and resolved. The prototype is ready for second-pass client review and final requirements sign-off before proceeding to stack selection and production development.")

add_body(doc, "This deliverable is not a production build. The current system is implemented entirely in vanilla JavaScript, HTML, and CSS with localStorage persistence. Its purpose is to replace static wireframes with interactive, data-backed screens that act as a living specification, enabling the client to validate business rules, user flows, and UI patterns before any stack or layer decisions are committed.", italic=True)


# ════════════════════════════════════════
#  SECTION 2: ALIGNMENT WITH PROPOSED TIMELINE
# ════════════════════════════════════════
add_heading(doc, "2)  Alignment with Proposed Timeline & Phases", level=2)

add_body(doc, "The original ERP Development Proposal (May 2026) outlines a 10–12 week delivery across six phases. Week 1 has completed Phases 1 and 2 in full, and has prototyped / pre-validated the core capabilities of Phases 3 and 4 as browser-level proof-of-concept. This accelerated progress means the project is on track to enter true Sprint 1 coding immediately upon client sign-off, with zero ambiguity on module scope or acceptance criteria.")

# Phase 1
add_heading(doc, "Phase 1: Discovery & Planning (Proposal Week 1) — ✅ COMPLETED", level=3)
add_bullet(doc, "Key Accomplishment:", "Instead of static documentation, the prototype itself is the requirements artifact. Stakeholders interact with real data, real gating rules, and real edge cases (cycle detection, role blocks, overdue aging) rather than reading about them.")

# Phase 2
add_heading(doc, "Phase 2: UI/UX Design & Prototyping (Proposal Weeks 1–2) — ✅ COMPLETED", level=3)
add_bullet(doc, "Key Accomplishment:", "The prototype went beyond wireframes by including responsive breakpoints, print media queries, empty-state patterns, and view-mode toggles (Table / Board / List). This means the Phase 3 frontend developer inherits a validated design system, not a blank canvas.")

# Phase 3
add_heading(doc, "Phase 3: Core Development — Sprint 1 (Proposal Weeks 2–5) — 🔄 PROTOTYPE PRE-VALIDATED", level=3)
add_body(doc, 'The proposal defines Sprint 1 as: "Authentication, user/role management, multi-entity client management, and workflow engine (task creation, assignment, dependencies, recurring templates). Backend APIs and frontend screens are developed in parallel. Deliverable: Deployed alpha on staging with core modules functional."', italic=True)

add_body(doc, "What the Prototype Already Demonstrates (Browser-Only Proof-of-Concept):", bold=True)
# Sprint 1 Table
add_table(doc,
    ["Sprint 1 Module", "Proposal Scope", "Prototype Pre-Validation"],
    [
        ["Authentication & RBAC",
         "JWT sessions, bcrypt hashing, permission matrix",
         "localStorage session with 4-tier role gates (Admin/Manager/Staff/Viewer); Auth.can() permission matrix enforced at module and action level; entity-scoped data visibility"],
        ["Multi-Entity Client Management",
         "CRUD, entity assignment, retainer status",
         "Full CRUD with mandatory ATA/LTA tag; related companies (parent, subsidiary, sister, affiliate); contact details multi-entry; point-of-contact assignment; retainer flag"],
        ["Workflow Engine",
         "Task creation, assignment, dependencies, recurring templates",
         "6-stage pipeline with gating rules; DAG predecessor selection with cycle detection; time logging; document upload metadata; retainer template generation; archive/restore"],
        ["Alpha Staging Readiness",
         "Deployed alpha environment",
         "The prototype is a self-contained static file set; it validates that all core user flows are coherent before a single server is provisioned"],
    ],
    col_widths=[Inches(1.8), Inches(2.0), Inches(3.0)]
)

add_body(doc, "Prototype Assumptions for Phase 3 (To Be Built in Production):", bold=True)
assumptions_p3 = [
    ("Real Backend:", "The prototype's localStorage DB engine will be replaced by a PostgreSQL database with row-level security policies for ATA/LTA segregation."),
    ("Real Authentication:", "The mock login (email/password stored in localStorage) will be replaced by JWT-based sessions with bcrypt hashing, short expiry, and refresh rotation."),
    ("Real APIs:", "All CRUD operations currently hitting localStorage will become RESTful API calls (Node.js/Django) with server-side validation."),
    ("Real Timezone Handling:", "The Manila-time-aware end-of-day reminder is client-side Intl.DateTimeFormat; production will use server-side UTC storage with Asia/Manila rendering."),
    ("Board View Interactivity:", "The prototype board view is read-only (no drag-and-drop). Production will implement drag-and-drop status changes with optimistic UI updates."),
    ("Risk Reduction:", "Because these modules were prototyped and client-tested in Week 1, the Sprint 1 development timeline can be compressed from 3 weeks to 2–2.5 weeks, with the remaining time allocated to automated testing and API hardening."),
]
for label, desc in assumptions_p3:
    add_bullet(doc, label, desc, indent_level=1)

# Phase 4
add_heading(doc, "Phase 4: Core Development — Sprint 2 (Proposal Weeks 5–8) — 🔄 PROTOTYPE PRE-VALIDATED", level=3)
add_body(doc, 'The proposal defines Sprint 2 as: "Billing module, disbursement module, document management system, and reporting dashboards. Integration between modules is completed and tested internally. Deliverable: Deployed beta with all MVP modules integrated."', italic=True)

add_body(doc, "What the Prototype Already Demonstrates (Browser-Only Proof-of-Concept):", bold=True)
# Sprint 2 Table
add_table(doc,
    ["Sprint 2 Module", "Proposal Scope", "Prototype Pre-Validation"],
    [
        ["Billing Module",
         "Invoice generation, PF/gov fee line items, payment recording, aging report, PDF export",
         "Sales Invoice with PF and Government Fee line items; auto-calculated subtotal; payment recording (Cash/Check/Bank Transfer/GCash/Maya); aging buckets (0–30, 31–60, 61–90, 90+); BIR-compliant PDF with seller header, TIN, and \"not valid for input tax\" footer; voucher print; trash/restore; recurring billing templates"],
        ["Disbursement Module",
         "Expense filing, approval workflow, client pass-through tracking, receipt upload",
         "Expense filing with category, amount, receipt metadata; 2-tier approval chain (Manager → Accounting for Firm Fund; 1-tier Accounting for Client Fund); \"Requested By\" and \"Payment Handled By\" tracking; disbursement PDF/voucher with signature lines; recurring disbursement templates"],
        ["Document Management System (DMS)",
         "File upload, categorization, original-copy tracking, version control, handover log",
         "Per-work-request upload with 4 categories; document lifecycle (collected → scanned → stored); envelope labeling; handover log with chain-of-custody; admin comments; view-mode toggles"],
        ["Reporting & Analytics",
         "Work request volume, task completion rate, billing summary, disbursement report, entity P&L snapshot",
         "Daily report (completions + time logs); weekly summary (completed/pending/overdue); monthly pending list with retainer awareness; revenue vs. disbursement bar chart; bento-grid KPI widgets; entity-scoped and consolidated views"],
        ["Module Integration",
         "Billing → Disbursement, Workflow → DMS, Reporting → all sources",
         "Routing blockers with dependency hints (e.g., \"Go to Billing module to create and link an invoice\"); linked disbursement IDs on invoices; DMS documents cross-referenced from task rows"],
    ],
    col_widths=[Inches(1.8), Inches(2.0), Inches(3.0)]
)

add_body(doc, "Prototype Assumptions for Phase 4 (To Be Built in Production):", bold=True)
assumptions_p4 = [
    ("Real File Storage:", "File uploads currently capture metadata only (filename, size, type) due to localStorage 5MB limits. Production will use AWS S3 / Cloudflare R2 for actual file persistence, presigned URLs, and encryption at rest."),
    ("Real PDF Generation:", "Current PDF output uses browser @media print CSS. Production will use a server-side PDF library (e.g., Puppeteer, WeasyPrint, or jsPDF on the server) for consistent rendering across devices."),
    ("Real Invoice Numbering:", "Sequential entity-specific numbering (ATA-SI-YYYY-NNN) is simulated. Production will use database sequences with atomic increments to prevent collisions."),
    ("Real Approval Workflows:", "Approval chains in the prototype are instant state changes. Production will implement async notification queues (email/Slack) and audit trails with timestamps and approver signatures."),
    ("Real Reporting Engine:", "Reports in the prototype aggregate localStorage arrays in memory. Production will use SQL aggregations with indexed queries for performance at scale."),
    ("Automated Recurring Generation:", "Retainer and billing templates require manual \"Generate Next Period\" clicks. Production will implement server-side cron jobs or scheduled task runners."),
    ("Full-Text Search:", "DMS search is metadata-only. Production will add PostgreSQL tsvector full-text search or a dedicated search index (e.g., Meilisearch) if document volume justifies it."),
    ("Risk Reduction:", "Because Sprint 2 modules were prototyped and client-tested in Week 1, the Sprint 2 timeline can also be compressed from 3 weeks to 2–2.5 weeks, with additional time for integration testing and security hardening."),
]
for label, desc in assumptions_p4:
    add_bullet(doc, label, desc, indent_level=1)

# Phase 5
add_heading(doc, "Phase 5: Integration Testing & UAT (Proposal Weeks 8–10) — ⏳ PLANNED", level=3)
add_bullet(doc, "Prototype Contribution to Phase 5:", "The 19 resolved feedback items from ERP Feedback v1 constitute an early UAT cycle. Issues like BIR PDF formatting, payment history overflow, and role clarity were caught and fixed before a single production commit. This de-risks the formal UAT window.")

# Phase 6
add_heading(doc, "Phase 6: Go-Live & Post-Launch Support (Proposal Weeks 10–12) — ⏳ PLANNED", level=3)


# ════════════════════════════════════════
#  SECTION 3: MODULE BREAKDOWN
# ════════════════════════════════════════
add_heading(doc, "3)  Week 1 Prototype Accomplishments — Module Breakdown", level=2)

# A) Foundation
add_heading(doc, "A)  Foundation & Data Layer (Browser-Only)", level=3)
modules_a = [
    ("localStorage DB Engine.", "Full CRUD persistence layer with schema versioning (v3), automatic migration from v2, and seed data hydration on first load."),
    ("Multi-Entity Segregation.", "Every record (clients, work requests, invoices, disbursements, documents) is tagged to ATA or LTA. The entity switcher in the header bar changes the active context and re-renders all views atomically."),
    ("Authentication & RBAC.", "Login with email/password, session restoration, four-tier role gates (Admin/Manager/Staff/Viewer), and permission matrix enforced at module and action level."),
    ("Schema v3 Design.", "Tables: users, clients, workRequests, tasks, invoices, disbursements, documents, transmittals, retainerTemplates, billingTemplates, disbursementTemplates, pendingChanges."),
]
for label, desc in modules_a:
    add_bullet(doc, label, desc)

# B) Dashboard
add_heading(doc, "B)  Dashboard & Analytics (Preview)", level=3)
modules_b = [
    ("Bento-Grid KPI Widgets.", "Active Work Requests, Total Clients, Disbursement Summary, and Revenue vs. Disbursement bar chart (SVG-based, mock data)."),
    ("Entity-Scoped vs. Consolidated Views.", "Admin sees combined ATA + LTA data; non-admin users see only their assigned entity."),
    ("Due-Date Warnings.", "Work Requests due this week are highlighted with color-coded urgency (Today = red, Tomorrow = orange)."),
]
for label, desc in modules_b:
    add_bullet(doc, label, desc)

# C) Client Management
add_heading(doc, "C)  Client Management Module", level=3)
modules_c = [
    ("Full CRUD.", "Create, edit, and archive clients with mandatory entity assignment, TIN, trade name, and business address."),
    ("Related Companies.", "Multi-entry UI for Parent, Subsidiary, Sister Company, and Affiliate relationships with bidirectional linking."),
    ("Contact Details.", "Repeatable rows for mobile, landline, email, and alternate contacts with labels."),
    ("Point of Contact.", "Employee selector assigning a staff member as the primary liaison for the client."),
    ("Retainer Flag.", "Checkbox marking the client as a recurring retainer; enables monthly/quarterly template generation."),
]
for label, desc in modules_c:
    add_bullet(doc, label, desc)

# D) Operations
add_heading(doc, "D)  Operations (Workflow & Task Management)", level=3)
modules_d = [
    ("Six-Stage Pipeline.", "Draft → Pre-processing → Processing → Billing → Disbursement → Completed. Each stage has gating rules validated before routing."),
    ("Task Dependency Engine (DAG).", "Predecessor selection per task; cycle detection prevents circular dependencies; downstream tasks block until upstream tasks are marked Completed."),
    ("Time Logging.", "Retrospective entry with start time, end time, auto-calculated hours (rounded to 0.25), and date. Inline per-task accordion in the detail view."),
    ("Document Upload.", "File metadata capture (filename, upload date, uploader ID) stored in taskDocuments[]; actual files are not persisted in localStorage to respect 5MB limits."),
    ("Retainer Templates.", "Manager-only feature: create reusable task templates with monthly or quarterly schedules; one-click generation clones tasks into a new work request."),
    ("Archive for Cancelled Work Requests.", "Cancelled items removed from active board/table views and moved to a dedicated Archive tab. All actions (time logs, uploads, status changes) are disabled on archived records. Restore to Draft available for managerial roles."),
]
for label, desc in modules_d:
    add_bullet(doc, label, desc)

# E) Billing
add_heading(doc, "E)  Billing Module", level=3)
modules_e = [
    ("Sales Invoice Creation.", "Line items for Professional Fee and Government Fee; auto-calculated subtotal; no VAT per Philippine EOPT Act compliance (total = subtotal)."),
    ("Payment Recording.", "Manual entry of amount, method (Cash/Check/Bank Transfer/GCash/Maya), reference number, date, and collected-by employee."),
    ("Aging Report.", "Overdue buckets: 0–30 days, 31–60 days, 61–90 days, 90+ days with color-coded totals."),
    ("BIR-Compliant Invoice PDF.", 'Browser print with seller header (name, TIN, address), buyer info, itemized list, totals, and footer: "This document is not valid for claim of input tax."'),
    ("Voucher Print.", "Stripped seller header for pass-through government fee receipts; minimal layout."),
    ("Trash & Restore.", "Draft invoices can be moved to Trash (Cancelled status); restored back to Draft."),
    ("Recurring Billing Templates.", 'Manual "Generate Next Period" button clones template into a new invoice.'),
]
for label, desc in modules_e:
    add_bullet(doc, label, desc)

# F) Disbursement
add_heading(doc, "F)  Disbursement & Expense Module", level=3)
modules_f = [
    ("Expense Filing.", "Staff submits category, description, amount, receipt (metadata), fund source (Firm Fund or Client Fund), and linked work request."),
    ("Approval Chain.", "1-tier (Admin) or 2-tier (Manager/Accounting) workflow: Submitted → Under Review → Approved → Released → Rejected."),
    ("Payment Handler Tracking.", 'Explicit "Requested By" and "Payment Handled By" fields on every disbursement record.'),
    ("Disbursement PDF / Voucher.", "Structured print with Voucher No, Date, Payee, Payment Method, Description, Amount, and signature lines for Prepared By + Approved By."),
    ("Recurring Templates.", "New Template form with Name, Category, Amount, Fund Source, and Schedule."),
]
for label, desc in modules_f:
    add_bullet(doc, label, desc)

# G) DMS
add_heading(doc, "G)  Document Management System (DMS)", level=3)
modules_g = [
    ("File Upload.", "Per-work-request upload with categorization (Requirement Docs, Processed Forms, Government Receipts, Final Deliverables)."),
    ("Document Lifecycle.", "Status transitions: collected → with_documentations → scanned → in_envelope → stored. Each transition is gated by Auth.can('dms:handover') permission."),
    ("Handover Log.", "Chain-of-custody entries for physical original copies handed to clients: handedOverTo, handoverDate, notes."),
    ("Envelope Labeling.", "Auto-generated text labels: ENVELOPE-[clientId]-[sequence]."),
    ("Admin Comments.", "Comment threads on document detail view; restricted to Admin role."),
]
for label, desc in modules_g:
    add_bullet(doc, label, desc)

# H) Transmittal
add_heading(doc, "H)  Transmittal Module", level=3)
modules_h = [
    ("Tracking Number Generation.", "Auto-sequential per entity: ATA-TX-2025-001, LTA-TX-2025-001."),
    ("Itemized Document List.", "Add/remove rows linking existing DMS documents; description and document type per item."),
    ("Acknowledgment.", '"Received By" text field + date entry when client physically signs the paper transmittal letter.'),
    ("Transmittal Letter PDF.", "Formal print with entity header, itemized list, signature block, and tracking number."),
]
for label, desc in modules_h:
    add_bullet(doc, label, desc)

# I) Reports
add_heading(doc, "I)  Reports Module (MVP-lite Preview)", level=3)
modules_i = [
    ("Daily Report.", "Tasks completed today per employee with start/end times and total hours."),
    ("Weekly Summary.", "Completed vs. pending vs. overdue per employee; overdue warnings highlighted."),
    ("Monthly Pending List.", "Pending tasks grouped by employee plus recurring retainer tasks due this month."),
]
for label, desc in modules_i:
    add_bullet(doc, label, desc)

# J) Admin
add_heading(doc, "J)  Admin / Users Module", level=3)
modules_j = [
    ("User CRUD.", "Create, edit, activate, and deactivate employee accounts with role assignment and entity scoping."),
    ("Pending Approvals (Admin Review Gate).", "Side-by-side diff of proposed changes vs. current record; Approve/Reject actions with rejection reason."),
    ("Self-Approval Block.", "Users cannot approve their own disbursements or review their own pending changes."),
]
for label, desc in modules_j:
    add_bullet(doc, label, desc)


# ════════════════════════════════════════
#  SECTION 4: CLIENT FEEDBACK INTEGRATION
# ════════════════════════════════════════
add_heading(doc, "4)  Client Feedback Integration — ERP Feedback v1", level=2)

add_body(doc, "The Week 1 prototype was delivered to the client for first-pass review on May 2026. The client returned 19 structured feedback items documented in ERP-feedback-v1.md. All items have been analyzed, triaged, and resolved in the prototype. A fix plan (ERP-feedback-v1-fix-plan.md) and implementation specification (ERP-feedback-v1-implementation-spec.md) were produced to ensure traceability.")

# A) Global UI
add_heading(doc, "A)  Global UI & Layout Fixes", level=3)
feedback_a = [
    ("Dashboard Widget Consistency.", "Issue #1: WR Due for Week missing from consolidated view. Resolved by adding the widget to both entity-scoped and consolidated dashboards."),
    ("Save/Cancel Button Placement.", "Issue #2: Buttons at bottom in Clients, Disbursement, Transmittal forms. Resolved by moving all Save/Cancel controls to the top-right form-header-bar pattern (mirroring Billing module)."),
    ("Board View Horizontal Scroll.", "Issues #9, #10: Billing and Operations board views overflowed horizontally. Resolved via proportional flex columns (flex:1; min-width:0), internal vertical scroll, and word-wrap on cards."),
    ("View-Mode Toggle Placement.", "Issue #12: Toggle sat in actions-bar instead of filters area. Resolved by relocating toggle below the filters bar in Disbursement and Transmittal modules."),
]
for label, desc in feedback_a:
    add_bullet(doc, label, desc)

# B) Admin Review Gate
add_heading(doc, "B)  Admin Review Gate Fixes", level=3)
add_bullet(doc, "Rejected Submissions Visibility.", "Issue #3: Staff rejected submissions disappeared. Resolved by adding a Rejected section in the Staff panel with rejection reason display and a Resubmit button that clones the snapshot into a new pending change.")

# C) Clients Module
add_heading(doc, "C)  Clients Module Fixes", level=3)
add_bullet(doc, "Missing Table Columns.", 'Issue #4: Related Companies and Contact Details not shown. Resolved by adding both columns to the clients table; rendering as comma-separated labels or "—" when empty.')

# D) Billing Module
add_heading(doc, "D)  Billing Module Fixes", level=3)
feedback_d = [
    ("Invoice PDF Format.", 'Issue #5: Missing BIR-compliant sections. Resolved by adding seller header (name, TIN, branch, address), buyer info, line items with qty/unit cost/total, subtotal = total, and footer: "This document is not valid for claim of input tax."'),
    ("Payment Details in List Views.", "Issue #6: Payment info absent from table/board/card. Resolved by adding Paid and Balance columns to the table view; balance displayed on board cards and list rows."),
    ("Voucher Print CSS.", "Issue #7: Overlay/popup styling minimal. Resolved with @media print CSS that hides nav, header, and actions; shows only invoice-detail content; BIR footer becomes block in print."),
    ("Payment History Overflow.", "Issue #8: Table overflows narrow container. Resolved by wrapping the payment history table in an overflow-x:auto container and condensing Recorded By + Collected By into a single Processor column."),
]
for label, desc in feedback_d:
    add_bullet(doc, label, desc)

# E) Disbursement Module
add_heading(doc, "E)  Disbursement Module Fixes", level=3)
feedback_e = [
    ("Role Clarity.", 'Issue #11: Requester vs handler unclear. Resolved by adding explicit "Requested By" and "Payment Handled By" labeled rows in the detail meta block.'),
    ("PDF Generation.", "Issue #13: Basic voucher format. Resolved with structured disbursement PDF including Voucher No, Date, Payee, Method, Description, Amount, and signature lines for Prepared By + Approved By."),
    ("Recurring Templates.", "Issue #14: No template creation UI. Resolved by adding a New Template form (Name, Category, Amount, Fund Source, Schedule, Description) and a Generate button."),
    ("Payment Info in List Views.", "Issue #6 (Disbursement part): Added Payment Method and Handled By columns in table view; method shown on board cards for Released items."),
]
for label, desc in feedback_e:
    add_bullet(doc, label, desc)

# F) Operations Module
add_heading(doc, "F)  Operations Module Fixes", level=3)
feedback_f = [
    ("Inline Time Logs.", "Issue #15: Time logs isolated in Task Activity section. Resolved by adding expandable accordion rows per task in the detail table; shows today's logs with a quick-add button."),
    ("Documentation Staff Visibility.", "Issue #16: Staff blocked from all non-assigned WRs. Resolved by adding Auth.can('dms:handover') exception; documentation staff sees all WRs in entity but cannot add/edit."),
    ("Month Filter Reliability.", "Issue #17: type=month input unreliable. Resolved by replacing with a select dropdown of 12 months + year, with explicit input and change event listeners."),
    ("Admin File Access.", "Issue #18: Admin could not view attached files. Resolved by adding clickable links that open DMS records in a new window; cross-linked when the file was uploaded via DMS."),
    ("Inline Task Documents.", "Issue #19: Documents isolated from task rows. Resolved by embedding document metadata, comments, and admin actions inside collapsible rows below each task."),
]
for label, desc in feedback_f:
    add_bullet(doc, label, desc)

# G) Additional Enhancements
add_heading(doc, "G)  Additional Enhancements Beyond Feedback v1", level=3)
feedback_g = [
    ("Widely Compatible View Icons.", "Centralized ViewIcons utility (Lucide-style SVGs) for Table, Board, and List toggles across all modules; ensures cross-browser compatibility."),
    ("Billing → Disbursement Routing Fix.", "Relaxed disbursement linkage gate to accept both task-level (linkedTaskId) and work-request-level (linkedWorkRequestId / linkedDisbursementIds) linked disbursements."),
    ("Actionable Dependency Hints.", 'Routing blocker messages now include italic hints (e.g., "→ Go to Billing module to create and link an invoice") with a Go button that navigates directly to the target module.'),
    ("End-of-Day Time Log Reminder.", "Manila-time-aware banner (5:00 PM+) prompting staff to log time for assigned tasks missing today's entry; includes a Log Time Now button opening the modal for the first missing task."),
    ("Routing Readiness Badges.", 'Board cards and list rows display "Ready to route" (green) or "N pending" (orange) badges with hover tooltips listing exact blockers.'),
    ("Status Color Consistency.", "Unified phase colors across progress bars, board columns, badges, route buttons, and the modern progress indicator."),
]
for label, desc in feedback_g:
    add_bullet(doc, label, desc)


# ════════════════════════════════════════
#  SECTION 5: KNOWN LIMITATIONS
# ════════════════════════════════════════
add_heading(doc, "5)  Known Limitations & Prototype Boundaries", level=2)

add_body(doc, "The following limitations are intentional for the requirements-finalization prototype. They will be addressed once the client approves the finalized specification and the project moves to stack selection and layer implementation.")

limitations = [
    ("No Server or Backend.", "All data is stored in browser localStorage. There is no REST API, no database, and no server-side validation. Data resets on browser cache clear unless exported."),
    ("No Real File Storage.", "File uploads capture metadata only (filename, size, type, uploader). Actual file content is not persisted to respect localStorage 5MB limits. Future stack will use AWS S3 / Cloudflare R2."),
    ("No Email or Notification Service.", "No SendGrid, Resend, or SMTP integration. All notifications are in-app badges and banners only."),
    ("No Mobile Native App.", "The prototype is web-responsive only. A native iOS/Android app is explicitly out of MVP scope per the proposal."),
    ("No Third-Party Integrations.", "No BIR eFPS, eBIRForms, QuickBooks, Xero, or bank API connections. These are Phase 2 / future enhancements."),
    ("Board View is Read-Only.", "Drag-and-drop status changes are not implemented. Cards are grouped by status for visibility only; status changes still use the existing dropdown/edit flow."),
    ("No Automated Recurring Generation.", 'Retainer and billing templates require manual "Generate Next Period" clicks. No server-side cron or scheduled jobs exist in a browser-only system.'),
]
for label, desc in limitations:
    add_bullet(doc, label, desc)


# ════════════════════════════════════════
#  SECTION 6: DELIVERABLES & FILE INVENTORY
# ════════════════════════════════════════
add_heading(doc, "6)  Deliverables & File Inventory", level=2)

add_body(doc, "The Week 1 prototype is delivered as a self-contained set of static files. No build step, no npm install, and no server is required. Opening index.html in any modern browser loads the full application.")

files = [
    ("index.html", "App shell, navigation structure, login screen, module script loading order, and responsive viewport meta."),
    ("css/styles.css", "Design system with CSS custom properties, responsive breakpoints, board view layout, print media queries, modal overlays, form patterns, badge styles, and bento-grid dashboard cards."),
    ("js/utils.js", "DOM builder (el()), PHP/date formatting, debounce, deepClone, generateId(), field error helpers, validateRequiredFields(), PaymentIcons map, and ViewIcons map."),
    ("js/data.js", "Schema v3 definition, seed data for all 12 tables, migration logic (v2 → v3), and the DB persistence engine (localStorage CRUD)."),
    ("js/auth.js", "Login/logout, session restore, role-based permission matrix (Auth.can()), entity switching, and user profile resolution."),
    ("js/app.js", "Hash router, module loader, sidebar highlight, view-mode preference persistence (localStorage), responsive menu toggle, and sidebar notification badges."),
    ("js/dashboard.js", "KPI widgets, revenue vs. disbursement chart, due-date warnings, upcoming disbursements, and entity-scoped/consolidated view rendering."),
    ("js/clients.js", "Client CRUD, related companies UI, contact details multi-entry, trade name, point-of-contact dropdown, retainer flag, and staff visibility filtering."),
    ("js/workflow.js", "Work request engine, 6-stage routing with gating, task DAG with cycle detection, time logging modal, document upload modal, inline expandable task rows, retainer templates, archive view, and end-of-day reminder."),
    ("js/billing.js", "Invoice CRUD, line items (PF/Gov Fee), payment recording, aging report, BIR invoice PDF, voucher print, trash/restore, recurring templates, and view-mode toggles."),
    ("js/disbursement.js", "Expense filing, 1-tier/2-tier approval chain, payment handler tracking, disbursement PDF/voucher, recurring templates, view-mode toggles, and pending-for-release section."),
    ("js/dms.js", "Document upload metadata, lifecycle transitions, handover log, envelope labeling, admin comments, view-mode toggles, and entity-scoped filtering."),
    ("js/transmittal.js", "Tracking number generation, itemized document list, acknowledgment form, transmittal letter PDF, and work request linkage."),
    ("js/reports.js", "Daily report (completions + time logs), weekly summary (completed/pending/overdue), monthly pending list with recurring retainer awareness, and filterable views."),
    ("js/users.js", "User CRUD, role/entity assignment, pending approvals tab with side-by-side diff, approve/reject actions, and self-approval block."),
    ("js/pendingChanges.js", "Admin Review Gate: submitForReview, approve, reject, resubmit, and buildDiff for side-by-side comparison."),
    ("js/datepicker.js + js/timepicker.js", "Material Design-inspired date and time picker components with overlay positioning and keyboard support."),
]
for fname, desc in files:
    add_bullet(doc, fname + "  —", desc)


# ════════════════════════════════════════
#  SECTION 7: NEXT STEPS
# ════════════════════════════════════════
add_heading(doc, "7)  Next Steps — Week 2 Plan", level=2)

add_body(doc, "Week 2 continues the requirements-finalization prototype cycle. The goal is to resolve any remaining feedback from ERP Feedback v1, add polish to high-touch UI elements, and produce a second prototype iteration (v1.1) ready for final client sign-off before stack selection.")

next_steps = [
    ("Client Feedback v1.1 Review.", "Address any new or unresolved feedback items from the client's review of the first round of fixes."),
    ("UI Polish Pass.", "Refine board card density, table row hover states, empty-state illustrations, and mobile breakpoint behavior."),
    ("Data Export / Import.", "Add JSON export/import for all tables so the client can preserve prototype data across sessions and share state with the development team."),
    ("Stack Decision Brief.", "Draft the technical architecture document (React vs. Vue, Node vs. Django, PostgreSQL schema, hosting platform) based on validated requirements."),
    ("UAT Test Script Draft.", "Begin writing the User Acceptance Testing checklist and test cases for Phase 5 (Proposal Weeks 8–10)."),
]
for label, desc in next_steps:
    add_bullet(doc, label, desc)


# ════════════════════════════════════════
#  SECTION 8: CONCLUSION & SIGN-OFF
# ════════════════════════════════════════
add_heading(doc, "8)  Conclusion & Sign-off", level=2)

add_body(doc, "Week 1 has successfully delivered a comprehensive, interactive browser-only prototype that validates the dual-entity accounting workflow, role-based access control, and all six operational stages from work request initiation through documentation handover.")

# Proposal deliverables tables
add_body(doc, "Proposal Deliverable Compliance:", bold=True)

# Table: Phase 1 Deliverables
add_table(doc,
    ["Proposal Deliverable", "Status", "How the Prototype Satisfies It"],
    [
        ["Requirements Specification", "✅ Complete", "12 modules, 80+ screens, and 19 resolved feedback items constitute a validated, living requirements spec"],
        ["Entity-Relationship Diagram (ERD)", "✅ Complete", "Schema v3 designed with 12 tables: users, clients, workRequests, tasks, invoices, disbursements, documents, transmittals, retainerTemplates, billingTemplates, disbursementTemplates, pendingChanges"],
        ["Sprint Backlog", "✅ Complete", "Each prototype screen maps to a user story; acceptance criteria are executable (the UI either works or doesn't)"],
        ["Business Rule Validation", "✅ Complete", "Domain rules (PF vs gov't fee, entity segregation, 2-tier approval, self-approval block) are enforced in code and tested against mock data"],
    ],
    col_widths=[Inches(1.8), Inches(1.2), Inches(3.8)]
)

add_table(doc,
    ["Proposal Deliverable", "Status", "How the Prototype Satisfies It"],
    [
        ["Clickable Prototype", "✅ Complete", "Fully navigable 12-module application with login, routing, modal flows, and print previews"],
        ["UI Style Guide", "✅ Complete", "CSS custom properties define the design system: color tokens, typography, spacing, form patterns, board layouts, badge styles, bento-grid dashboard"],
        ["Design Approval Readiness", "✅ Complete", "Client received the prototype, submitted 19 feedback items, and all items were triaged and resolved"],
    ],
    col_widths=[Inches(1.8), Inches(1.2), Inches(3.8)]
)

# Testing & Quality Assurance
add_body(doc, "Testing & Quality Assurance:", bold=True)
add_table(doc,
    ["Proposal Deliverable", "Status", "Notes"],
    [
        ["Unit Testing (≥ 80% coverage)", "⏳ Planned", "Will begin in Sprint 1 (Week 2) per proposal; prototype acceptance criteria serve as test cases"],
        ["Integration Testing", "⏳ Planned", "Critical workflow end-to-end tests: WR creation → task assignment → billing → document upload"],
        ["UAT Sign-off", "⏳ Planned", "UAT test script draft is a Week 2 deliverable (see Next Steps); Domain Expert (CPA/MBA) will validate financial outputs"],
        ["Security Testing", "⏳ Planned", "OWASP ZAP baseline scan, RBAC policy validation, dependency vulnerability scanning"],
    ],
    col_widths=[Inches(1.8), Inches(1.2), Inches(3.8)]
)

# Go-Live & Support
add_body(doc, "Go-Live & Support:", bold=True)
add_table(doc,
    ["Proposal Deliverable", "Status", "Notes"],
    [
        ["Production Deployment", "⏳ Planned", "Hosting platform (Vercel / Railway / Render) to be selected after client sign-off"],
        ["User Training", "⏳ Planned", "The prototype itself serves as a training sandbox; users can practice without touching production data"],
        ["Hypercare (14 days)", "⏳ Planned", "Critical bug resolution within 12 hours; non-critical within 48 hours"],
        ["Documentation", "⏳ Planned", "System administrator manual and end-user quick-start guide"],
    ],
    col_widths=[Inches(1.8), Inches(1.2), Inches(3.8)]
)

# Phase Status Summary
add_body(doc, "Phase Status Summary:", bold=True)
add_table(doc,
    ["Phase", "Proposal Weeks", "Status"],
    [
        ["Phase 1: Discovery & Planning", "Week 1", "✅ COMPLETED — Requirements spec, ERD, and sprint backlog validated through prototype"],
        ["Phase 2: UI/UX Design & Prototyping", "Weeks 1–2", "✅ COMPLETED — Clickable prototype delivered; UI style guide established; 19 feedback items resolved"],
        ["Phase 3: Sprint 1 — Core Engine", "Weeks 2–5", "🔄 PROTOTYPE PRE-VALIDATED — Auth, roles, client mgmt, workflow engine demonstrated in browser; assumptions defined for production build"],
        ["Phase 4: Sprint 2 — Business Modules", "Weeks 5–8", "🔄 PROTOTYPE PRE-VALIDATED — Billing, disbursement, DMS, reporting demonstrated in browser; assumptions defined for production build"],
        ["Phase 5: Testing & UAT", "Weeks 8–10", "⏳ PLANNED — UAT script draft in Week 2; formal UAT scheduled after production beta"],
        ["Phase 6: Go-Live & Support", "Weeks 10–12", "⏳ PLANNED — Deployment, training, and 14-day hypercare"],
    ],
    col_widths=[Inches(2.2), Inches(1.2), Inches(3.4)]
)

# Project Timeline Impact
add_body(doc, "Project Timeline Impact:", bold=True)
add_body(doc, "Because Phases 1–2 are complete and Phases 3–4 capabilities are pre-validated through the prototype, the project is positioned to begin production coding immediately upon client sign-off, with compressed sprint timelines and significantly reduced scope-ambiguity risk. The prototype itself serves as the single source of truth for acceptance criteria, UI patterns, and business rule enforcement.")

add_body(doc, "This report is submitted for client review and written approval. Upon sign-off, the project will transition from the Requirements Finalization phase to Month 1 (Milestone 1) — Platform Foundation production development.", italic=True, space_after=Pt(24))


# ════════════════════════════════════════
#  SIGN-OFF BLOCK
# ════════════════════════════════════════

# Divider
p_div = doc.add_paragraph()
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_div, "─" * 75, size=Pt(8), color=RGBColor(0xBF, 0xBF, 0xBF))
p_div.paragraph_format.space_after = Pt(12)

# Sign-off table (no header shading — clean signature block)
sign_table = doc.add_table(rows=3, cols=2)
sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(sign_table)

# Row 0 - Labels
for ci, label in enumerate(["Prepared By:", "Reviewed By:"]):
    cell = sign_table.rows[0].cells[ci]
    set_cell_shading(cell, "F2F2F2")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, label, bold=True, size=Pt(11))

# Row 1 - Names
cell_prep = sign_table.rows[1].cells[0]
p_prep = cell_prep.paragraphs[0]
add_run(p_prep, "Jayvee Marcelo", bold=True, size=Pt(11))
p_prep2 = cell_prep.add_paragraph()
add_run(p_prep2, "Full-Stack Developer / Prototype Lead", size=SMALL_SIZE)
p_prep3 = cell_prep.add_paragraph()
add_run(p_prep3, "MicroAxis Solutions Corp.", size=SMALL_SIZE)

cell_rev = sign_table.rows[1].cells[1]
p_rev = cell_rev.paragraphs[0]
p_rev.paragraph_format.space_before = Pt(24)
add_run(p_rev, "___________________________________________________", size=Pt(11))
p_rev2 = cell_rev.add_paragraph()
add_run(p_rev2, "Client Representative", size=SMALL_SIZE)
p_rev3 = cell_rev.add_paragraph()
add_run(p_rev3, "Date: ___________________________________________________", size=SMALL_SIZE)

# Row 2 - Date
cell_date = sign_table.rows[2].cells[0]
p_date = cell_date.paragraphs[0]
add_run(p_date, "Date: June 6, 2026", size=SMALL_SIZE)

sign_table.rows[2].cells[1].paragraphs[0].text = ""

doc.add_paragraph()  # spacing

# End of Report
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p_end, "— End of Report —", bold=True, size=Pt(12))


# ── Save ──
output_path = "Week_1_Accomplishments_Report_v2.docx"
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
